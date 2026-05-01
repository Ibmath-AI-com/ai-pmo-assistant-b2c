from __future__ import annotations

import re
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(rendered_content: str, output_path: str) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in rendered_content.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(text, style="List Number")
        else:
            p = doc.add_paragraph()
            _apply_inline_formatting(p, line)

    doc.save(output_path)
    return output_path


def _apply_inline_formatting(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def create_temp_docx(rendered_content: str) -> str:
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    generate_docx(rendered_content, tmp.name)
    return tmp.name
